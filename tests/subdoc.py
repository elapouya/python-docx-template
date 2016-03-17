# -*- coding: utf-8 -*-
'''
Created : 2015-03-12

@author: Eric Lapouyade
'''

from docxtpl import DocxTemplate
from docx.shared import Inches

tpl=DocxTemplate('test_files/subdoc_tpl.docx')

sd = tpl.new_subdoc()
p = sd.add_paragraph('This is a sub-document inserted into a bigger one')
p = sd.add_paragraph('It has been ')
p.add_run('dynamically').style = 'dynamic'
p.add_run(' generated with python by using ')
p.add_run('python-docx').italic = True
p.add_run(' library')

sd.add_heading('Heading, level 1', level=1)
sd.add_paragraph('This is an Intense quote', style='IntenseQuote')

sd.add_paragraph('A picture :')
sd.add_picture('test_files/python_logo.png', width=Inches(1.25))

sd.add_paragraph('A Table :')
table = sd.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
recordset=( (1,101,'Spam'),
            (2,42,'Eggs'),
            (3,631,'Spam,spam, eggs, and ham') )
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item[0])
    row_cells[1].text = str(item[1])
    row_cells[2].text = item[2]

context = {
    'mysubdoc' : sd,
}

tpl.render(context)
tpl.save('test_files/subdoc.docx')
