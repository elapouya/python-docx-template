# -*- coding: utf-8 -*-
"""
Created : 2015-03-12

@author: Eric Lapouyade
"""
from __future__ import annotations

import binascii
import functools
import io
import os
import re
import zipfile
from os import PathLike
from typing import IO, TYPE_CHECKING, Any, Generator

import docx.oxml.ns
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as REL_TYPE
from docx.opc.oxml import parse_xml
from docx.opc.part import XmlPart
from jinja2 import Environment, Template, meta
from jinja2.exceptions import TemplateError
from lxml import etree

from .subdoc import Subdoc

if TYPE_CHECKING:
    from docx.document import Document as _Document
    from docx.oxml.document import CT_Document
    from docx.parts.story import StoryPart

    class DocumentObject(_Document):
        _element: CT_Document


class DocxTemplate:
    """Class for managing docx files as they were jinja2 templates"""

    HEADER_URI = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
    )
    FOOTER_URI = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
    )

    def __init__(self, template_file: IO[bytes] | str | PathLike) -> None:
        self.template_file = template_file
        self.reset_replacements()
        self.docx: DocumentObject = None  # type:ignore[assignment]
        self.is_rendered = False
        self.is_saved = False
        self.allow_missing_pics = False

    def init_docx(self, reload: bool = True) -> None:
        if not self.docx or (self.is_rendered and reload):
            self.docx = Document(self.template_file)  # type:ignore[arg-type,assignment]
            self.is_rendered = False

    def render_init(self) -> None:
        self.init_docx()
        self.pic_map: dict = {}
        self.current_rendering_part: StoryPart = None  # type:ignore[assignment]
        self.docx_ids_index = 1000
        self.is_saved = False

    def __getattr__(self, name) -> Any:
        return getattr(self.docx, name)

    def xml_to_string(self, xml, encoding="unicode") -> str:
        # Be careful : pretty_print MUST be set to False, otherwise patch_xml()
        # won't work properly
        return etree.tostring(xml, encoding="unicode", pretty_print=False)

    def get_docx(self) -> DocumentObject:
        self.init_docx()
        return self.docx

    def get_xml(self) -> str:
        return self.xml_to_string(self.docx._element.body)

    def write_xml(self, filename: str | PathLike) -> None:
        with open(filename, "w") as fh:
            fh.write(self.get_xml())

    def patch_xml(self, src_xml: str) -> str:
        """Make a lots of cleaning to have a raw xml understandable by jinja2 :
        strip all unnecessary xml tags, manage table cell background color and colspan,
        unescape html entities, etc..."""

        # replace {<something>{ by {{   ( works with {{ }} {% and %} {# and #})
        src_xml = re.sub(
            r"(?<={)(<[^>]*>)+(?=[\{%\#])|(?<=[%\}\#])(<[^>]*>)+(?=\})",
            "",
            src_xml,
            flags=re.DOTALL,
        )

        # replace {{<some tags>jinja2 stuff<some other tags>}} by {{jinja2 stuff}}
        # same thing with {% ... %} and {# #}
        # "jinja2 stuff" could a variable, a 'if' etc... anything jinja2 will understand
        def striptags(m) -> str:
            return re.sub(
                "</w:t>.*?(<w:t>|<w:t [^>]*>)", "", m.group(0), flags=re.DOTALL
            )

        src_xml = re.sub(
            r"{%(?:(?!%}).)*|{#(?:(?!#}).)*|{{(?:(?!}}).)*",
            striptags,
            src_xml,
            flags=re.DOTALL,
        )

        # manage table cell colspan
        def colspan(m) -> str:
            cell_xml = m.group(1) + m.group(3)
            cell_xml = re.sub(
                r"<w:r[ >](?:(?!<w:r[ >]).)*<w:t></w:t>.*?</w:r>",
                "",
                cell_xml,
                flags=re.DOTALL,
            )
            cell_xml = re.sub(r"<w:gridSpan[^/]*/>", "", cell_xml, count=1)
            return re.sub(
                r"(<w:tcPr[^>]*>)",
                r'\1<w:gridSpan w:val="{{%s}}"/>' % m.group(2),
                cell_xml,
            )

        src_xml = re.sub(
            r"(<w:tc[ >](?:(?!<w:tc[ >]).)*){%\s*colspan\s+([^%]*)\s*%}(.*?</w:tc>)",
            colspan,
            src_xml,
            flags=re.DOTALL,
        )

        # manage table cell background color
        def cellbg(m) -> str:
            cell_xml = m.group(1) + m.group(3)
            cell_xml = re.sub(
                r"<w:r[ >](?:(?!<w:r[ >]).)*<w:t></w:t>.*?</w:r>",
                "",
                cell_xml,
                flags=re.DOTALL,
            )
            cell_xml = re.sub(r"<w:shd[^/]*/>", "", cell_xml, count=1)
            return re.sub(
                r"(<w:tcPr[^>]*>)",
                r'\1<w:shd w:val="clear" w:color="auto" w:fill="{{%s}}"/>' % m.group(2),
                cell_xml,
            )

        src_xml = re.sub(
            r"(<w:tc[ >](?:(?!<w:tc[ >]).)*){%\s*cellbg\s+([^%]*)\s*%}(.*?</w:tc>)",
            cellbg,
            src_xml,
            flags=re.DOTALL,
        )

        # ensure space preservation
        src_xml = re.sub(
            r"<w:t>((?:(?!<w:t>).)*)({{.*?}}|{%.*?%})",
            r'<w:t xml:space="preserve">\1\2',
            src_xml,
            flags=re.DOTALL,
        )
        src_xml = re.sub(
            r"({{r\s.*?}}|{%r\s.*?%})",
            r'</w:t></w:r><w:r><w:t xml:space="preserve">\1</w:t></w:r><w:r><w:t xml:space="preserve">',
            src_xml,
            flags=re.DOTALL,
        )

        # {%- will merge with previous paragraph text
        src_xml = re.sub(r"</w:t>(?:(?!</w:t>).)*?{%-", "{%", src_xml, flags=re.DOTALL)
        # -%} will merge with next paragraph text
        src_xml = re.sub(
            r"-%}(?:(?!<w:t[ >]|{%|{{).)*?<w:t[^>]*?>", "%}", src_xml, flags=re.DOTALL
        )

        for y in ["tr", "tc", "p", "r"]:
            # replace into xml code the row/paragraph/run containing
            # {%y xxx %} or {{y xxx}} template tag
            # by {% xxx %} or {{ xx }} without any surrounding <w:y> tags :
            # This is mandatory to have jinja2 generating correct xml code
            pat = (
                r"<w:%(y)s[ >](?:(?!<w:%(y)s[ >]).)*({%%|{{)%(y)s ([^}%%]*(?:%%}|}})).*?</w:%(y)s>"
                % {"y": y}
            )
            src_xml = re.sub(pat, r"\1 \2", src_xml, flags=re.DOTALL)

        for y in ["tr", "tc", "p"]:
            # same thing, but for {#y xxx #} (but not where y == 'r', since that
            # makes less sense to use comments in that context
            pat = (
                r"<w:%(y)s[ >](?:(?!<w:%(y)s[ >]).)*({#)%(y)s ([^}#]*(?:#})).*?</w:%(y)s>"
                % {"y": y}
            )
            src_xml = re.sub(pat, r"\1 \2", src_xml, flags=re.DOTALL)

        # add vMerge
        # use {% vm %} to make this table cell and its copies
        # be vertically merged within a {% for %}
        def v_merge_tc(m) -> str:
            def v_merge(m1) -> str:
                return (
                    '<w:vMerge w:val="{% if loop.first %}restart{% else %}continue{% endif %}"/>'
                    + m1.group(1)  # Everything between ``</w:tcPr>`` and ``<w:t>``.
                    + "{% if loop.first %}"
                    + m1.group(2)  # Everything before ``{% vm %}``.
                    + m1.group(3)  # Everything after ``{% vm %}``.
                    + "{% endif %}"
                    + m1.group(4)  # ``</w:t>``.
                )

            return re.sub(
                r"(</w:tcPr[ >].*?<w:t(?:.*?)>)(.*?)(?:{%\s*vm\s*%})(.*?)(</w:t>)",
                v_merge,
                m.group(),
                # Everything between ``</w:tc>`` and ``</w:tc>`` with ``{% vm %}`` inside.
                flags=re.DOTALL,
            )

        src_xml = re.sub(
            r"<w:tc[ >](?:(?!<w:tc[ >]).)*?{%\s*vm\s*%}.*?</w:tc[ >]",
            v_merge_tc,
            src_xml,
            flags=re.DOTALL,
        )

        # Use ``{% hm %}`` to make table cell become horizontally merged within
        # a ``{% for %}``.
        def h_merge_tc(m) -> str:
            xml_to_patch = (
                m.group()
            )  # Everything between ``</w:tc>`` and ``</w:tc>`` with ``{% hm %}`` inside.

            def with_gridspan(m1) -> str:
                return (
                    m1.group(1)  # ``w:gridSpan w:val="``.
                    + "{{ "
                    + m1.group(2)
                    + " * loop.length }}"  # Content of ``w:val``, multiplied by loop length.
                    + m1.group(3)  # Closing quotation mark.
                )

            def without_gridspan(m2) -> str:
                return (
                    '<w:gridSpan w:val="{{ loop.length }}"/>'
                    + m2.group(1)  # Everything between ``</w:tcPr>`` and ``<w:t>``.
                    + m2.group(2)  # Everything before ``{% hm %}``.
                    + m2.group(3)  # Everything after ``{% hm %}``.
                    + m2.group(4)  # ``</w:t>``.
                )

            if re.search(r"w:gridSpan", xml_to_patch):
                # Simple case, there's already ``gridSpan``, multiply its value.

                xml = re.sub(
                    r'(w:gridSpan w:val=")(\d+)(")',
                    with_gridspan,
                    xml_to_patch,
                    flags=re.DOTALL,
                )
                xml = re.sub(
                    r"{%\s*hm\s*%}",
                    "",
                    xml,  # Patched xml.
                    flags=re.DOTALL,
                )
            else:
                # There're no ``gridSpan``, add one.
                xml = re.sub(
                    r"(</w:tcPr[ >].*?<w:t(?:.*?)>)(.*?)(?:{%\s*hm\s*%})(.*?)(</w:t>)",
                    without_gridspan,
                    xml_to_patch,
                    flags=re.DOTALL,
                )

            # Discard every other cell generated in loop.
            return "{% if loop.first %}" + xml + "{% endif %}"

        src_xml = re.sub(
            r"<w:tc[ >](?:(?!<w:tc[ >]).)*?{%\s*hm\s*%}.*?</w:tc[ >]",
            h_merge_tc,
            src_xml,
            flags=re.DOTALL,
        )

        def clean_tags(m) -> str:
            return (
                m.group(0)
                .replace(r"&#8216;", "'")
                .replace("&lt;", "<")
                .replace("&gt;", ">")
                .replace("“", '"')
                .replace("”", '"')
                .replace("‘", "'")
                .replace("’", "'")
            )

        src_xml = re.sub(r"(?<=\{[\{%])(.*?)(?=[\}%]})", clean_tags, src_xml)

        return src_xml

    def render_xml_part(
        self, src_xml: str, part: StoryPart, context: dict[str, Any], jinja_env=None
    ) -> str:
        src_xml = re.sub(r"<w:p([ >])", r"\n<w:p\1", src_xml)
        try:
            self.current_rendering_part = part
            template = (
                jinja_env.from_string(src_xml) if jinja_env else Template(src_xml)
            )
            dst_xml = template.render(context)
        except TemplateError as exc:
            if hasattr(exc, "lineno") and exc.lineno is not None:
                line_number = max(exc.lineno - 4, 0)
                exc.docx_context = map(  # type:ignore[attr-defined]
                    lambda x: re.sub(r"<[^>]+>", "", x),
                    src_xml.splitlines()[line_number: (line_number + 7)],  # fmt: skip
                )

            raise exc
        dst_xml = re.sub(r"\n<w:p([ >])", r"<w:p\1", dst_xml)
        dst_xml = (
            dst_xml.replace("{_{", "{{")
            .replace("}_}", "}}")
            .replace("{_%", "{%")
            .replace("%_}", "%}")
        )
        dst_xml = self.resolve_listing(dst_xml)
        return dst_xml

    def render_properties(
        self, context: dict[str, Any], jinja_env: Environment | None = None
    ) -> None:
        # List of string attributes of docx.opc.coreprops.CoreProperties which are strings.
        # It seems that some attributes cannot be written as strings. Those are commented out.
        properties = [
            "author",
            # 'category',
            "comments",
            # 'content_status',
            "identifier",
            # 'keywords',
            "language",
            # 'last_modified_by',
            "subject",
            "title",
            # 'version',
        ]
        if jinja_env is None:
            jinja_env = Environment()

        for prop in properties:
            initial = getattr(self.docx.core_properties, prop)
            template = jinja_env.from_string(initial)
            rendered = template.render(context)
            setattr(self.docx.core_properties, prop, rendered)

    def render_footnotes(
        self, context: dict[str, Any], jinja_env: Environment | None = None
    ) -> None:
        if jinja_env is None:
            jinja_env = Environment()

        for section in self.docx.sections:
            for part in section.part.package.parts:
                if part.content_type == (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.footnotes+xml"
                ):
                    xml = self.patch_xml(
                        part.blob.decode("utf-8")
                        if isinstance(part.blob, bytes)
                        else part.blob
                    )
                    xml = self.render_xml_part(xml, part, context, jinja_env)
                    part._blob = xml.encode("utf-8")

    def resolve_listing(self, xml) -> str:
        def resolve_text(run_properties, paragraph_properties, m) -> str:
            xml = m.group(0).replace(
                "\t",
                "</w:t></w:r>"
                "<w:r>%s<w:tab/></w:r>"
                '<w:r>%s<w:t xml:space="preserve">' % (run_properties, run_properties),
            )
            xml = xml.replace(
                "\a",
                "</w:t></w:r></w:p>"
                '<w:p>%s<w:r>%s<w:t xml:space="preserve">'
                % (paragraph_properties, run_properties),
            )
            xml = xml.replace("\n", '</w:t><w:br/><w:t xml:space="preserve">')
            xml = xml.replace(
                "\f",
                "</w:t></w:r></w:p>"
                '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'
                '<w:p>%s<w:r>%s<w:t xml:space="preserve">'
                % (paragraph_properties, run_properties),
            )
            return xml

        def resolve_run(paragraph_properties, m) -> str:
            m_run_properties = re.search(r"<w:rPr>.*?</w:rPr>", m.group(0))
            run_properties = m_run_properties.group(0) if m_run_properties else ""
            return re.sub(
                r"<w:t(?: [^>]*)?>.*?</w:t>",
                lambda x: resolve_text(run_properties, paragraph_properties, x),
                m.group(0),
                flags=re.DOTALL,
            )

        def resolve_paragraph(m) -> str:
            m_paragraph_properties = re.search(r"<w:pPr>.*?</w:pPr>", m.group(0))
            paragraph_properties = (
                m_paragraph_properties.group(0) if m_paragraph_properties else ""
            )
            return re.sub(
                r"<w:r(?: [^>]*)?>.*?</w:r>",
                lambda x: resolve_run(paragraph_properties, x),
                m.group(0),
                flags=re.DOTALL,
            )

        xml = re.sub(
            r"<w:p(?: [^>]*)?>.*?</w:p>", resolve_paragraph, xml, flags=re.DOTALL
        )

        return xml

    def build_xml(self, context, jinja_env=None) -> str:
        xml = self.get_xml()
        xml = self.patch_xml(xml)
        xml = self.render_xml_part(xml, self.docx._part, context, jinja_env)
        return xml

    def map_tree(self, tree) -> None:
        root = self.docx._element
        body = root.body
        root.replace(body, tree)

    def get_headers_footers(self, uri: str) -> Generator[tuple[str, StoryPart]]:
        for relKey, val in self.docx._part.rels.items():
            if (val.reltype == uri) and (val.target_part.blob):
                yield relKey, val.target_part

    def get_part_xml(self, part) -> str:
        return self.xml_to_string(parse_xml(part.blob))

    def get_headers_footers_encoding(self, xml) -> str:
        m = re.match(r'<\?xml[^\?]+\bencoding="([^"]+)"', xml, re.I)
        if m:
            return m.group(1)
        return "utf-8"

    def build_headers_footers_xml(
        self, context, uri, jinja_env=None
    ) -> Generator[tuple[str, bytes]]:
        for relKey, part in self.get_headers_footers(uri):
            xml = self.get_part_xml(part)
            encoding = self.get_headers_footers_encoding(xml)
            xml = self.patch_xml(xml)
            xml = self.render_xml_part(xml, part, context, jinja_env)
            yield relKey, xml.encode(encoding)

    def map_headers_footers_xml(self, relKey, xml) -> None:
        part = self.docx._part.rels[relKey].target_part
        new_part = XmlPart.load(part.partname, part.content_type, xml, part.package)
        for rId, rel in part.rels.items():
            new_part.load_rel(rel.reltype, rel._target, rel.rId, rel.is_external)
        self.docx._part.rels[relKey]._target = new_part

    def render(
        self,
        context: dict[str, Any],
        jinja_env: Environment | None = None,
        autoescape: bool = False,
    ) -> None:
        # init template working attributes
        self.render_init()

        if autoescape:
            if not jinja_env:
                jinja_env = Environment(autoescape=autoescape)
            else:
                jinja_env.autoescape = autoescape

        # Body
        xml_src = self.build_xml(context, jinja_env)

        # fix tables if needed
        tree = self.fix_tables(xml_src)

        # fix docPr ID's
        self.fix_docpr_ids(tree)

        # Replace body xml tree
        self.map_tree(tree)

        # Headers
        headers = self.build_headers_footers_xml(context, self.HEADER_URI, jinja_env)
        for relKey, xml in headers:
            self.map_headers_footers_xml(relKey, xml)

        # Footers
        footers = self.build_headers_footers_xml(context, self.FOOTER_URI, jinja_env)
        for relKey, xml in footers:
            self.map_headers_footers_xml(relKey, xml)

        self.render_properties(context, jinja_env)

        self.render_footnotes(context, jinja_env)

        # set rendered flag
        self.is_rendered = True

    # using of TC tag in for cycle can cause that count of columns does not
    # correspond to real count of columns in row. This function is able to fix it.
    def fix_tables(self, xml: str) -> etree._Element:
        parser = etree.XMLParser(recover=True)
        tree = etree.fromstring(xml, parser=parser)
        # get namespace
        ns = "{" + tree.nsmap["w"] + "}"
        # walk trough xml and find table
        for t in tree.iter(ns + "tbl"):
            tblGrid = t.find(ns + "tblGrid")
            if tblGrid is None:
                continue
            columns = tblGrid.findall(ns + "gridCol")
            to_add = 0
            # walk trough all rows and try to find if there is higher cell count
            for r in t.iter(ns + "tr"):
                cells = r.findall(ns + "tc")
                if (len(columns) + to_add) < len(cells):
                    to_add = len(cells) - len(columns)
            # is necessary to add columns?
            if to_add > 0:
                # at first, calculate width of table according to columns
                # (we want to preserve it)
                width = 0.0
                new_average = None
                ns_w = ns + "w"
                for c in columns:
                    w = c.get(ns_w)
                    if w is not None:
                        width += float(w)
                # try to keep proportion of table
                if width > 0:
                    old_average = width / len(columns)
                    new_average = width / (len(columns) + to_add)
                    # scale the old columns
                    for c in columns:
                        w = c.get(ns_w)
                        if w is not None:
                            c.set(ns_w, str(int(float(w) * new_average / old_average)))
                    # add new columns
                    for i in range(to_add):
                        etree.SubElement(
                            tblGrid, ns + "gridCol", {ns_w: str(int(new_average))}
                        )

            # Refetch columns after columns addition.
            columns = tblGrid.findall(ns + "gridCol")
            columns_len = len(columns)

            cells_len_max = 0

            def get_cell_len(total: int, cell) -> int:
                tc_pr = cell.find(ns + "tcPr")
                grid_span = None if tc_pr is None else tc_pr.find(ns + "gridSpan")

                if grid_span is not None:
                    return total + int(grid_span.get(ns + "val"))

                return total + 1

            # Calculate max of table cells to compare with `gridCol`.
            for r in t.iter(ns + "tr"):
                cells = r.findall(ns + "tc")
                cells_len = functools.reduce(get_cell_len, cells, 0)
                cells_len_max = max(cells_len_max, cells_len)

            to_remove = columns_len - cells_len_max

            # If after the loop, there're less columns, than
            # originally was, remove extra `gridCol` declarations.
            if to_remove > 0:
                # Have to keep track of the removed width to scale the
                # table back to its original width.
                removed_width = 0.0

                ns_w = ns + "w"
                for c in columns[-to_remove:]:
                    w = c.get(ns_w)
                    if w is not None:
                        removed_width += float(w)

                    tblGrid.remove(c)

                columns_left = tblGrid.findall(ns + "gridCol")

                # Distribute `removed_width` across all columns that has
                # left after extras removal.
                extra_space = 0
                if len(columns_left) > 0:
                    extra_space = int(removed_width / len(columns_left))

                for c in columns_left:
                    w = c.get(ns_w)
                    if w is not None:
                        c.set(ns_w, str(int(float(w)) + extra_space))

        return tree

    def fix_docpr_ids(self, tree) -> None:
        # some Ids may have some collisions : so renumbering all of them :
        for elt in tree.xpath("//wp:docPr", namespaces=docx.oxml.ns.nsmap):
            self.docx_ids_index += 1
            elt.attrib["id"] = str(self.docx_ids_index)

    def new_subdoc(self, docpath=None) -> Subdoc:
        self.init_docx()
        return Subdoc(self, docpath)

    @staticmethod
    def get_file_crc(file_obj: IO[bytes] | str | PathLike) -> int:
        if hasattr(file_obj, "read"):
            buf = file_obj.read()
        else:
            with open(file_obj, "rb") as fh:
                buf = fh.read()

        crc = binascii.crc32(buf) & 0xFFFFFFFF
        return crc

    def replace_media(self, src_file, dst_file) -> None:
        """Replace one media by another one into a docx

        This has been done mainly because it is not possible to add images in
        docx header/footer.
        With this function, put a dummy picture in your header/footer,
        then specify it with its replacement in this function using the file path
        or file-like objects.

        Syntax: tpl.replace_media('dummy_media_to_replace.png','media_to_paste.jpg')
            -- or --
                tpl.replace_media(io.BytesIO(image_stream), io.BytesIO(new_image_stream))

        Note: for images, the aspect ratio will be the same as the replaced image

        Note2: it is important to have the source media file as it is required
                to calculate its CRC to find them in the docx
        """

        crc = self.get_file_crc(src_file)
        if hasattr(dst_file, "read"):
            self.crc_to_new_media[crc] = dst_file.read()
        else:
            with open(dst_file, "rb") as fh:
                self.crc_to_new_media[crc] = fh.read()

    def replace_pic(self, embedded_file, dst_file) -> None:
        """Replace embedded picture with original-name given by embedded_file.
           (give only the file basename, not the full path)
           The new picture is given by dst_file (either a filename or a file-like
           object)

        Notes:
            1) embedded_file and dst_file must have the same extension/format
               in case dst_file is a file-like object, no check is done on
               format compatibility
            2) the aspect ratio will be the same as the replaced image
            3) There is no need to keep the original file (this is not the case
               for replace_embedded and replace_media)
        """

        if hasattr(dst_file, "read"):
            # NOTE: file extension not checked
            self.pics_to_replace[embedded_file] = dst_file.read()
        else:
            with open(dst_file, "rb") as fh:
                self.pics_to_replace[embedded_file] = fh.read()

    def replace_embedded(self, src_file, dst_file) -> None:
        """Replace one embedded object by another one into a docx

        This has been done mainly because it is not possible to add images
        in docx header/footer.
        With this function, put a dummy picture in your header/footer,
        then specify it with its replacement in this function

        Syntax: tpl.replace_embedded('dummy_doc.docx','doc_to_paste.docx')

        Note2 : it is important to have the source file as it is required to
                calculate its CRC to find them in the docx
        """
        with open(dst_file, "rb") as fh:
            crc = self.get_file_crc(src_file)
            self.crc_to_new_embedded[crc] = fh.read()

    def replace_zipname(self, zipname, dst_file) -> None:
        """Replace one file in the docx file

        First note that a MSWord .docx file is in fact a zip file.

        This method can be used to replace document embedded in the docx template.

        Some embedded document may have been modified by MSWord while saving
        the template : thus replace_embedded() cannot be used as CRC is not the
        same as the original file.

        This method works for embedded MSWord file like Excel or PowerPoint file,
        but won't work for others like PDF, Python or even Text files :
        For these ones, MSWord generate an oleObjectNNN.bin file which is no
        use to be replaced as it is encoded.

        Syntax:

        tpl.replace_zipname(
            'word/embeddings/Feuille_Microsoft_Office_Excel1.xlsx',
            'my_excel_file.xlsx')

        The zipname is the one you can find when you open docx with WinZip,
        7zip (Windows) or unzip -l (Linux). The zipname starts with
        "word/embeddings/". Note that the file is renamed by MSWord,
        so you have to guess a little bit...
        """
        with open(dst_file, "rb") as fh:
            self.zipname_to_replace[zipname] = fh.read()

    def reset_replacements(self) -> None:
        """Reset replacement dictionaries

        This will reset data for image/embedded/zipname replacement

        This is useful when calling several times render() with different
        image/embedded/zipname replacements without re-instantiating
        DocxTemplate object.
        In this case, the right sequence for each rendering will be :
            - reset_replacements(...)
            - replace_zipname(...), replace_media(...) and/or replace_embedded(...),
            - render(...)

        If you instantiate DocxTemplate object before each render(),
        this method is useless.
        """
        self.crc_to_new_media: dict = {}
        self.crc_to_new_embedded: dict = {}
        self.zipname_to_replace: dict = {}
        self.pics_to_replace: dict = {}

    def post_processing(self, docx_file) -> None:
        if self.crc_to_new_media or self.crc_to_new_embedded or self.zipname_to_replace:
            if hasattr(docx_file, "read"):
                bytes_io = io.BytesIO()
                DocxTemplate(docx_file).save(bytes_io)
                bytes_io.seek(0)
                docx_file.seek(0)
                docx_file.truncate()
                docx_file.seek(0)
                tmp_file: IO[bytes] | str = bytes_io

            else:
                tmp_file = "%s_docxtpl_before_replace_medias" % docx_file
                os.rename(docx_file, tmp_file)

            with zipfile.ZipFile(tmp_file) as zin:
                with zipfile.ZipFile(docx_file, "w") as zout:
                    for item in zin.infolist():
                        buf = zin.read(item.filename)
                        if item.filename in self.zipname_to_replace:
                            zout.writestr(item, self.zipname_to_replace[item.filename])
                        elif (
                            item.filename.startswith("word/media/")
                            and item.CRC in self.crc_to_new_media
                        ):
                            zout.writestr(item, self.crc_to_new_media[item.CRC])
                        elif (
                            item.filename.startswith("word/embeddings/")
                            and item.CRC in self.crc_to_new_embedded
                        ):
                            zout.writestr(item, self.crc_to_new_embedded[item.CRC])
                        else:
                            zout.writestr(item, buf)

            if not hasattr(tmp_file, "read"):
                os.remove(tmp_file)
            if hasattr(docx_file, "read"):
                docx_file.seek(0)

    def pre_processing(self) -> None:
        if self.pics_to_replace:
            self._replace_pics()

    def _replace_pics(self) -> None:
        """Replaces pictures xml tags in the docx template with pictures provided by the user"""

        replaced_pics = {key: False for key in self.pics_to_replace}

        # Main document
        part = self.docx.part
        self._replace_docx_part_pics(part, replaced_pics)

        # Header/Footer
        for relid, rel in part.rels.items():
            if rel.reltype in (REL_TYPE.HEADER, REL_TYPE.FOOTER):
                self._replace_docx_part_pics(rel.target_part, replaced_pics)

        if not self.allow_missing_pics:
            # make sure all template images defined by user were replaced
            for img_id, replaced in replaced_pics.items():
                if not replaced:
                    raise ValueError(
                        "Picture %s not found in the docx template" % img_id
                    )

    def get_pic_map(self) -> dict:
        return self.pic_map

    def _replace_docx_part_pics(self, doc_part, replaced_pics) -> None:
        et = etree.fromstring(doc_part.blob)

        part_map: dict = {}

        gds = et.xpath("//a:graphic/a:graphicData", namespaces=docx.oxml.ns.nsmap)
        for gd in gds:  # type:ignore[union-attr]
            rel = None
            # Either IMAGE, CHART, SMART_ART, ...
            if not hasattr(gd, "attrib"):
                continue
            try:
                uri = gd.attrib["uri"]
            except KeyError:
                continue
            try:
                if uri == docx.oxml.ns.nsmap["pic"]:
                    # Either PICTURE or LINKED_PICTURE image
                    blip = gd.xpath(  # type:ignore[union-attr,index]
                        "pic:pic/pic:blipFill/a:blip", namespaces=docx.oxml.ns.nsmap
                    )[0]
                    dest = blip.xpath(  # type:ignore[union-attr]
                        "@r:embed", namespaces=docx.oxml.ns.nsmap
                    )
                    try:
                        rel = dest[0]  # type:ignore[index]
                    except (IndexError, KeyError, TypeError):
                        continue
                else:
                    continue

                non_visual_properties = "pic:pic/pic:nvPicPr/pic:cNvPr/"
                filename = gd.xpath(  # type:ignore[union-attr,index]
                    "%s@name" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )[0]
                titles = gd.xpath(  # type:ignore[union-attr]
                    "%s@title" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )
                title = titles[0] if titles else ""  # type:ignore[index]
                descriptions = gd.xpath(  # type:ignore[union-attr]
                    "%s@descr" % non_visual_properties, namespaces=docx.oxml.ns.nsmap
                )
                description = (
                    descriptions[0] if descriptions else ""  # type:ignore[index]
                )

                part_map[filename] = (
                    doc_part.rels[rel].target_ref,
                    doc_part.rels[rel].target_part,
                )

                # replace data
                for img_id, img_data in self.pics_to_replace.items():
                    if img_id in (filename, title, description):
                        part_map[filename][1]._blob = img_data
                        replaced_pics[img_id] = True
                        break

            # FIXME: figure out what exceptions are thrown here
            # and catch more specific exceptions
            except Exception:
                continue

        self.pic_map.update(part_map)

    def build_url_id(self, url):
        self.init_docx()
        return self.docx._part.relate_to(url, REL_TYPE.HYPERLINK, is_external=True)

    def save(self, filename: IO[bytes] | str | PathLike, *args, **kwargs) -> None:
        # case where save() is called without doing rendering
        # ( user wants only to replace image/embedded/zipname )
        if not self.is_saved and not self.is_rendered:
            self.docx = Document(self.template_file)  # type:ignore[arg-type,assignment]
        self.pre_processing()
        self.docx.save(filename, *args, **kwargs)  # type:ignore[arg-type]
        self.post_processing(filename)
        self.is_saved = True

    def get_undeclared_template_variables(
        self,
        jinja_env: Environment | None = None,
        context: dict[str, Any] | None = None,
    ) -> set[str]:
        # Create a temporary document to analyze the template without affecting the current state
        temp_doc = Document(self.template_file)  # type:ignore[arg-type]

        # Get XML from the temporary document
        xml = self.xml_to_string(temp_doc._element.body)  # type:ignore[attr-defined]
        xml = self.patch_xml(xml)

        # Add headers and footers
        for uri in [self.HEADER_URI, self.FOOTER_URI]:
            for relKey, val in temp_doc._part.rels.items():
                if (val.reltype == uri) and (val.target_part.blob):
                    _xml = self.xml_to_string(parse_xml(val.target_part.blob))
                    xml += self.patch_xml(_xml)

        env = jinja_env or Environment()

        parse_content = env.parse(xml)
        all_variables = meta.find_undeclared_variables(parse_content)

        # If context is provided, return only variables that are not in the context
        if context is not None:
            provided_variables = set(context.keys())
            return all_variables - provided_variables

        # If no context provided, return all variables (original behavior)
        return all_variables
